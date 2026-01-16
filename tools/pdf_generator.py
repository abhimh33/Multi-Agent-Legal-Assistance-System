# tools/pdf_generator.py
"""
Professional PDF Generator for Legal Documents.
Provides comprehensive PDF generation with proper formatting, headers, footers,
and professional styling suitable for legal documents.
"""

import os
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
from io import BytesIO

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm, mm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, ListFlowable, ListItem, KeepTogether
)
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from dotenv import load_dotenv

load_dotenv()


class LegalDocumentStyles:
    """Predefined styles for legal documents."""
    
    @staticmethod
    def get_styles() -> Dict[str, ParagraphStyle]:
        """Get all document styles."""
        base_styles = getSampleStyleSheet()
        
        styles = {
            'title': ParagraphStyle(
                'Title',
                parent=base_styles['Title'],
                fontSize=18,
                fontName='Helvetica-Bold',
                alignment=TA_CENTER,
                spaceAfter=20,
                spaceBefore=10,
                textColor=colors.HexColor('#1a365d')
            ),
            'subtitle': ParagraphStyle(
                'Subtitle',
                parent=base_styles['Normal'],
                fontSize=14,
                fontName='Helvetica-Bold',
                alignment=TA_CENTER,
                spaceAfter=15,
                spaceBefore=10,
                textColor=colors.HexColor('#2d3748')
            ),
            'heading1': ParagraphStyle(
                'Heading1',
                parent=base_styles['Heading1'],
                fontSize=14,
                fontName='Helvetica-Bold',
                spaceBefore=15,
                spaceAfter=8,
                textColor=colors.HexColor('#1a365d'),
                borderWidth=0,
                borderColor=colors.HexColor('#e2e8f0'),
                borderPadding=5
            ),
            'heading2': ParagraphStyle(
                'Heading2',
                parent=base_styles['Heading2'],
                fontSize=12,
                fontName='Helvetica-Bold',
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#2d3748')
            ),
            'body': ParagraphStyle(
                'Body',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                spaceBefore=6,
                spaceAfter=6,
                leading=14,
                firstLineIndent=0
            ),
            'body_indented': ParagraphStyle(
                'BodyIndented',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                spaceBefore=4,
                spaceAfter=4,
                leading=14,
                leftIndent=20,
                firstLineIndent=0
            ),
            'clause': ParagraphStyle(
                'Clause',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica',
                alignment=TA_JUSTIFY,
                spaceBefore=8,
                spaceAfter=4,
                leading=14,
                leftIndent=30,
                bulletIndent=15
            ),
            'signature': ParagraphStyle(
                'Signature',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica',
                alignment=TA_LEFT,
                spaceBefore=30,
                spaceAfter=5
            ),
            'footer': ParagraphStyle(
                'Footer',
                parent=base_styles['Normal'],
                fontSize=8,
                fontName='Helvetica-Oblique',
                alignment=TA_CENTER,
                textColor=colors.grey
            ),
            'disclaimer': ParagraphStyle(
                'Disclaimer',
                parent=base_styles['Normal'],
                fontSize=9,
                fontName='Helvetica-Oblique',
                alignment=TA_CENTER,
                textColor=colors.HexColor('#718096'),
                spaceBefore=20,
                spaceAfter=10
            ),
            'bold': ParagraphStyle(
                'Bold',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT
            ),
            'centered': ParagraphStyle(
                'Centered',
                parent=base_styles['Normal'],
                fontSize=11,
                fontName='Helvetica',
                alignment=TA_CENTER
            ),
        }
        
        return styles


class HeaderFooter:
    """Custom header and footer for PDF documents."""
    
    def __init__(self, title: str = "Legal Document", include_page_numbers: bool = True):
        self.title = title
        self.include_page_numbers = include_page_numbers
    
    def __call__(self, canvas: canvas.Canvas, doc):
        """Add header and footer to each page."""
        canvas.saveState()
        
        # Header
        canvas.setFont('Helvetica-Bold', 10)
        canvas.setFillColor(colors.HexColor('#1a365d'))
        canvas.drawString(doc.leftMargin, doc.height + doc.topMargin - 10, self.title)
        
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(
            doc.width + doc.leftMargin,
            doc.height + doc.topMargin - 10,
            datetime.now().strftime('%B %d, %Y')
        )
        
        # Header line
        canvas.setStrokeColor(colors.HexColor('#e2e8f0'))
        canvas.setLineWidth(0.5)
        canvas.line(
            doc.leftMargin,
            doc.height + doc.topMargin - 20,
            doc.width + doc.leftMargin,
            doc.height + doc.topMargin - 20
        )
        
        # Footer
        if self.include_page_numbers:
            canvas.setFont('Helvetica', 9)
            canvas.setFillColor(colors.grey)
            page_text = f"Page {doc.page}"
            canvas.drawCentredString(doc.width / 2 + doc.leftMargin, 25, page_text)
        
        # Footer line
        canvas.line(
            doc.leftMargin,
            40,
            doc.width + doc.leftMargin,
            40
        )
        
        # Footer text
        canvas.setFont('Helvetica-Oblique', 7)
        canvas.drawCentredString(
            doc.width / 2 + doc.leftMargin,
            15,
            "Generated by AI Legal Assistant - For Informational Purposes Only"
        )
        
        canvas.restoreState()


class PDFGenerator:
    """Professional PDF generator for legal documents."""
    
    def __init__(
        self,
        page_size: Tuple = A4,
        margins: Tuple[float, float, float, float] = (1*inch, 1*inch, 0.75*inch, 0.75*inch)
    ):
        """
        Initialize PDF generator.
        
        Args:
            page_size: Page size (A4 or letter)
            margins: (left, right, top, bottom) margins in inches
        """
        self.page_size = page_size
        self.margins = margins  # left, right, top, bottom
        self.styles = LegalDocumentStyles.get_styles()
        self.export_dir = os.getenv("EXPORT_DIRECTORY", "exports")
    
    def generate_pdf(
        self,
        content: str,
        filename: str = None,
        title: str = "Legal Document",
        include_header_footer: bool = True,
        return_bytes: bool = False
    ) -> Dict[str, Any]:
        """
        Generate a PDF from text content.
        
        Args:
            content: The document content (plain text or structured)
            filename: Output filename (without extension)
            title: Document title for header
            include_header_footer: Whether to include header/footer
            return_bytes: If True, return PDF as bytes instead of saving
            
        Returns:
            Dict with status and file path or bytes
        """
        try:
            # Create exports directory
            if not return_bytes:
                Path(self.export_dir).mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"legal_document_{timestamp}"
            
            filename = filename.replace(" ", "_").replace("/", "_")
            
            # Create PDF
            if return_bytes:
                buffer = BytesIO()
                output = buffer
            else:
                filepath = os.path.join(self.export_dir, f"{filename}.pdf")
                output = filepath
            
            doc = SimpleDocTemplate(
                output,
                pagesize=self.page_size,
                leftMargin=self.margins[0],
                rightMargin=self.margins[1],
                topMargin=self.margins[2] + 30,  # Extra space for header
                bottomMargin=self.margins[3] + 30  # Extra space for footer
            )
            
            # Build story (content elements)
            story = self._parse_content_to_story(content, title)
            
            # Build PDF with header/footer
            if include_header_footer:
                header_footer = HeaderFooter(title=title)
                doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
            else:
                doc.build(story)
            
            if return_bytes:
                buffer.seek(0)
                return {
                    "status": "success",
                    "data": buffer.getvalue(),
                    "filename": f"{filename}.pdf"
                }
            else:
                return {
                    "status": "success",
                    "filepath": filepath,
                    "filename": f"{filename}.pdf",
                    "message": f"PDF generated successfully: {filepath}"
                }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"PDF generation failed: {str(e)}"
            }
    
    def _parse_content_to_story(self, content: str, title: str) -> List:
        """Parse text content into ReportLab story elements."""
        story = []
        
        # Add title
        story.append(Paragraph(title.upper(), self.styles['title']))
        story.append(Spacer(1, 20))
        
        # Parse content line by line
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            
            # Detect section headers (lines with === or ---)
            if stripped.startswith('=' * 5) or stripped.startswith('-' * 5):
                story.append(HRFlowable(
                    width="100%",
                    thickness=1,
                    color=colors.HexColor('#e2e8f0'),
                    spaceBefore=10,
                    spaceAfter=10
                ))
                continue
            
            # Detect headings (all caps or numbered sections)
            if stripped.isupper() and len(stripped) > 3:
                story.append(Paragraph(stripped, self.styles['heading1']))
                current_section = stripped
                continue
            
            # Detect numbered clauses
            if self._is_numbered_clause(stripped):
                story.append(Paragraph(stripped, self.styles['body']))
                continue
            
            # Detect sub-items (a), b), etc.)
            if self._is_sub_item(stripped):
                story.append(Paragraph(stripped, self.styles['body_indented']))
                continue
            
            # Regular paragraph
            # Escape special characters for ReportLab
            safe_text = self._escape_text(stripped)
            story.append(Paragraph(safe_text, self.styles['body']))
        
        # Add disclaimer
        story.append(Spacer(1, 30))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Paragraph(
            "DISCLAIMER: This document is generated for informational purposes only. "
            "Please consult a qualified legal professional before using this document "
            "for any legal proceedings or official purposes.",
            self.styles['disclaimer']
        ))
        
        return story
    
    def _is_numbered_clause(self, text: str) -> bool:
        """Check if text is a numbered clause."""
        import re
        patterns = [
            r'^\d+\.',  # 1. 2. 3.
            r'^\d+\)',  # 1) 2) 3)
            r'^[IVX]+\.',  # I. II. III.
            r'^[ivx]+\.',  # i. ii. iii.
        ]
        return any(re.match(p, text) for p in patterns)
    
    def _is_sub_item(self, text: str) -> bool:
        """Check if text is a sub-item."""
        import re
        patterns = [
            r'^[a-z]\)',  # a) b) c)
            r'^[a-z]\.',  # a. b. c.
            r'^\([a-z]\)',  # (a) (b) (c)
            r'^[•\-\*]',  # bullet points
        ]
        return any(re.match(p, text) for p in patterns)
    
    def _escape_text(self, text: str) -> str:
        """Escape special characters for ReportLab XML."""
        replacements = [
            ('&', '&amp;'),
            ('<', '&lt;'),
            ('>', '&gt;'),
        ]
        for old, new in replacements:
            text = text.replace(old, new)
        return text
    
    def generate_from_template(
        self,
        template_content: str,
        filename: str = None,
        title: str = None
    ) -> Dict[str, Any]:
        """
        Generate PDF from a template-generated document.
        
        Args:
            template_content: Content from a DocumentTemplate.generate()
            filename: Output filename
            title: Document title (extracted from content if not provided)
            
        Returns:
            Dict with status and file path
        """
        # Try to extract title from content
        if title is None:
            lines = template_content.split('\n')
            for line in lines:
                stripped = line.strip()
                if stripped and not stripped.startswith('='):
                    title = stripped
                    break
            title = title or "Legal Document"
        
        return self.generate_pdf(
            content=template_content,
            filename=filename,
            title=title,
            include_header_footer=True
        )


class DocumentExporter:
    """Unified document exporter for multiple formats."""
    
    def __init__(self):
        self.pdf_generator = PDFGenerator()
        self.export_dir = os.getenv("EXPORT_DIRECTORY", "exports")
        Path(self.export_dir).mkdir(parents=True, exist_ok=True)
    
    def export(
        self,
        content: str,
        format_type: str = "pdf",
        filename: str = None,
        title: str = "Legal Document"
    ) -> Dict[str, Any]:
        """
        Export document to specified format.
        
        Args:
            content: Document content
            format_type: 'pdf', 'txt', 'docx', or 'all'
            filename: Base filename (without extension)
            title: Document title
            
        Returns:
            Dict with export results
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"legal_document_{timestamp}"
        
        results = {}
        
        # Export as TXT
        if format_type.lower() in ["txt", "text", "all"]:
            results["txt"] = self._export_txt(content, filename)
        
        # Export as PDF
        if format_type.lower() in ["pdf", "all"]:
            results["pdf"] = self.pdf_generator.generate_pdf(
                content=content,
                filename=filename,
                title=title
            )
        
        # Export as DOCX
        if format_type.lower() in ["docx", "word", "all"]:
            results["docx"] = self._export_docx(content, filename, title)
        
        # Single format
        if format_type.lower() == "pdf":
            return results.get("pdf", {"status": "error", "message": "PDF export failed"})
        elif format_type.lower() in ["txt", "text"]:
            return results.get("txt", {"status": "error", "message": "TXT export failed"})
        elif format_type.lower() in ["docx", "word"]:
            return results.get("docx", {"status": "error", "message": "DOCX export failed"})
        
        return {
            "status": "success",
            "exports": results,
            "export_directory": self.export_dir
        }
    
    def _export_txt(self, content: str, filename: str) -> Dict[str, Any]:
        """Export to plain text file."""
        try:
            filepath = os.path.join(self.export_dir, f"{filename}.txt")
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            return {
                "status": "success",
                "filepath": filepath,
                "message": f"Text file saved: {filepath}"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def _export_docx(self, content: str, filename: str, title: str) -> Dict[str, Any]:
        """Export to DOCX file."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            lines = content.split('\n')
            for line in lines:
                stripped = line.strip()
                
                if not stripped:
                    doc.add_paragraph()
                    continue
                
                if stripped.startswith('=' * 5) or stripped.startswith('-' * 5):
                    continue
                
                if stripped.isupper() and len(stripped) > 3:
                    doc.add_heading(stripped, level=1)
                else:
                    para = doc.add_paragraph(stripped)
                    para.paragraph_format.space_after = Pt(6)
            
            filepath = os.path.join(self.export_dir, f"{filename}.docx")
            doc.save(filepath)
            
            return {
                "status": "success",
                "filepath": filepath,
                "message": f"DOCX file saved: {filepath}"
            }
        except ImportError:
            return {
                "status": "error",
                "message": "python-docx not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def get_pdf_bytes(self, content: str, title: str = "Legal Document") -> bytes:
        """Get PDF as bytes for streaming/download."""
        result = self.pdf_generator.generate_pdf(
            content=content,
            title=title,
            return_bytes=True
        )
        if result["status"] == "success":
            return result["data"]
        return None


# Convenience function for CrewAI tool
def generate_legal_pdf(content: str, filename: str = None, title: str = "Legal Document") -> Dict[str, Any]:
    """Generate a professional legal PDF document."""
    exporter = DocumentExporter()
    return exporter.export(content, format_type="pdf", filename=filename, title=title)
